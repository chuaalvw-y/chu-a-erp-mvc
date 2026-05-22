#!/usr/bin/env python3
"""
Build a single Word (.docx) and PDF rendering of the ChuA.ERP end-user
documentation set from the Markdown sources in this folder.

Pipeline (pure-Python, no system dependencies):
    Markdown --(markdown lib)--> HTML --(BeautifulSoup walk)--> DOCX + PDF

Usage:
    python docs/build-docs.py

Outputs (next to this script):
    ChuA.ERP-User-Documentation.docx
    ChuA.ERP-User-Documentation.pdf

Notes:
- Mermaid code fences are rendered as monospaced "diagram source" boxes,
  since no diagram-rendering engine (mermaid-cli / browser) is available
  in this environment.
- The DOCX Table of Contents is inserted as a Word field; open the file
  in Word and choose "Update field" (F9) to populate page numbers.
- The PDF Table of Contents is generated with live page numbers via a
  two-pass reportlab build.
"""

import datetime as _dt
import html as _html
import os
import re
import sys

import markdown as _md
from bs4 import BeautifulSoup, NavigableString, Tag

# ----------------------------------------------------------------------------
# Source files, in reading order (matches docs/README.md site map).
# ----------------------------------------------------------------------------
HERE = os.path.dirname(os.path.abspath(__file__))
DOC_FILES = [
    "README.md",
    "user-guide/01-introduction.md",
    "user-guide/02-getting-started.md",
    "user-guide/03-logging-in.md",
    "user-guide/04-navigation.md",
    "user-guide/05-dashboard.md",
    "user-guide/06-user-profile.md",
    "user-guide/07-notifications.md",
    "user-guide/08-workflow-approvals.md",
    "user-guide/09-reports-exports.md",
    "user-guide/10-search-filtering.md",
    "user-guide/11-attachments.md",
    "user-guide/12-audit-history.md",
    "modules/purchasing.md",
    "modules/sales.md",
    "modules/inventory.md",
    "modules/finance.md",
    "modules/crm.md",
    "modules/workflow-engine.md",
    "admin/administration.md",
    "admin/system-settings.md",
    "admin/multi-tenant.md",
    "admin/security-permissions.md",
    "reference/quick-start.md",
    "reference/troubleshooting.md",
    "reference/faq.md",
    "reference/glossary.md",
    "reference/role-based-guides.md",
]

TITLE = "ChuA.ERP"
SUBTITLE = "End-User Documentation"
VERSION = "Phase 5 — MVC UI (v1 API)"
OUT_BASE = "ChuA.ERP-User-Documentation"

MD_EXTENSIONS = ["tables", "fenced_code", "sane_lists", "attr_list"]


def md_to_soup(md_text: str) -> BeautifulSoup:
    html_text = _md.markdown(md_text, extensions=MD_EXTENSIONS)
    return BeautifulSoup(html_text, "html.parser")


def clean_text(node) -> str:
    return node.get_text()


def wrap_pre_lines(text: str, width: int = 95) -> str:
    """Soft-wrap very long preformatted lines so they don't overflow the PDF frame."""
    out_lines = []
    for line in text.split("\n"):
        while len(line) > width:
            out_lines.append(line[:width])
            line = line[width:]
        out_lines.append(line)
    return "\n".join(out_lines)


# ============================================================================
# DOCX rendering
# ============================================================================
import docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_background(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _add_inline_runs(paragraph, node, bold=False, italic=False, mono=False, link=False):
    """Recursively add runs to a docx paragraph from a BS4 inline node."""
    if isinstance(node, NavigableString):
        run = paragraph.add_run(str(node))
        run.bold = bold or None
        run.italic = italic or None
        if mono:
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        if link:
            run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)
            run.font.underline = True
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    nb = bold or name in ("strong", "b")
    ni = italic or name in ("em", "i")
    nm = mono or name == "code"
    nl = link or name == "a"
    for child in node.children:
        _add_inline_runs(paragraph, child, nb, ni, nm, nl)


def _docx_add_code_block(doc, text):
    text = text.rstrip("\n")
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _set_cell_background(cell, "F4F4F4")
    cell.paragraphs[0].text = ""
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    # subtle border
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "CCCCCC")
        borders.append(e)
    tbl_pr.append(borders)


def _docx_add_callout(doc, node):
    """Render a blockquote as an indented callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # collect inline content from inner paragraphs
    inner = node.find_all("p")
    if not inner:
        _add_inline_runs(p, node)
    else:
        for i, para in enumerate(inner):
            if i > 0:
                p.add_run("\n")
            _add_inline_runs(p, para)
    for run in p.runs:
        if run.font.color and run.font.color.rgb is None:
            run.font.color.rgb = RGBColor(0x49, 0x50, 0x57)


def _docx_add_list(doc, node, level=0):
    ordered = node.name.lower() == "ol"
    style = "List Number" if ordered else "List Bullet"
    if level > 0:
        style = f"{style} {min(level + 1, 3)}"
    for li in node.find_all("li", recursive=False):
        p = doc.add_paragraph(style=style)
        # inline content of li (excluding nested lists)
        for child in li.children:
            if isinstance(child, Tag) and child.name in ("ul", "ol"):
                continue
            _add_inline_runs(p, child)
        # nested lists
        for child in li.find_all(["ul", "ol"], recursive=False):
            _docx_add_list(doc, child, level + 1)


def _docx_add_table(doc, node):
    rows = node.find_all("tr")
    if not rows:
        return
    ncols = max(len(r.find_all(["td", "th"])) for r in rows)
    table = doc.add_table(rows=0, cols=ncols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, tr in enumerate(rows):
        cells = tr.find_all(["td", "th"])
        row_cells = table.add_row().cells
        for c_idx in range(ncols):
            target = row_cells[c_idx]
            target.paragraphs[0].text = ""
            if c_idx < len(cells):
                para = target.paragraphs[0]
                _add_inline_runs(para, cells[c_idx])
                if cells[c_idx].name == "th":
                    for run in para.runs:
                        run.bold = True


def _docx_render_block(doc, node):
    if isinstance(node, NavigableString):
        if str(node).strip():
            doc.add_paragraph(str(node).strip())
        return
    if not isinstance(node, Tag):
        return
    name = node.name.lower()
    if re.fullmatch(r"h[1-6]", name):
        level = int(name[1])
        heading = doc.add_heading(level=min(level, 4))
        _add_inline_runs(heading, node)
    elif name == "p":
        p = doc.add_paragraph()
        _add_inline_runs(p, node)
    elif name in ("ul", "ol"):
        _docx_add_list(doc, node)
    elif name == "table":
        _docx_add_table(doc, node)
        doc.add_paragraph()
    elif name == "pre":
        code = node.find("code")
        _docx_add_code_block(doc, code.get_text() if code else node.get_text())
    elif name == "blockquote":
        _docx_add_callout(doc, node)
    elif name == "hr":
        doc.add_paragraph("_" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # fall back: render children
        for child in node.children:
            _docx_render_block(doc, child)


def _docx_add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = r'TOC \o "1-3" \h \z \u'
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose “Update Field” to build the table of contents."
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(t)
    run._r.append(fldChar3)


def build_docx(out_path):
    doc = Document()
    # python-docx writes <w:zoom w:val="bestFit"/> which omits the schema-required
    # w:percent attribute. Add it so the file passes strict OOXML validation.
    _zoom = doc.settings.element.find(qn("w:zoom"))
    if _zoom is not None:
        _zoom.set(qn("w:percent"), "100")
    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(36)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(SUBTITLE)
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x49, 0x50, 0x57)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"{VERSION}\nGenerated {_dt.date.today().isoformat()}")
    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    _docx_add_toc(doc)
    doc.add_page_break()

    for rel in DOC_FILES:
        path = os.path.join(HERE, rel)
        if not os.path.exists(path):
            continue
        with open(path, encoding="utf-8") as f:
            soup = md_to_soup(f.read())
        for block in list(soup.children):
            _docx_render_block(doc, block)
        doc.add_page_break()

    doc.save(out_path)
    print(f"  DOCX written: {out_path}")


# ============================================================================
# PDF rendering (reportlab)
# ============================================================================
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import (
    BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, PageBreak,
    Table, TableStyle, Preformatted, ListFlowable, ListItem, HRFlowable,
    KeepTogether,
)
from reportlab.platypus.tableofcontents import TableOfContents


def _inline_to_rl(node) -> str:
    """Convert a BS4 inline node to reportlab mini-markup."""
    if isinstance(node, NavigableString):
        return _html.escape(str(node))
    if not isinstance(node, Tag):
        return ""
    name = node.name.lower()
    inner = "".join(_inline_to_rl(c) for c in node.children)
    if name in ("strong", "b"):
        return f"<b>{inner}</b>"
    if name in ("em", "i"):
        return f"<i>{inner}</i>"
    if name == "code":
        return f'<font face="Courier" size="9">{inner}</font>'
    if name == "a":
        href = node.get("href", "")
        if href.startswith("http"):
            return f'<a href="{_html.escape(href)}" color="#0d6efd">{inner}</a>'
        return f'<font color="#0d6efd">{inner}</font>'
    if name == "br":
        return "<br/>"
    return inner


class ErpDocTemplate(BaseDocTemplate):
    def __init__(self, filename, **kw):
        super().__init__(filename, **kw)
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="main")
        self.addPageTemplates([PageTemplate(id="body", frames=[frame], onPage=self._footer)])
        self._toc_entries = []

    def _footer(self, canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.grey)
        canvas.drawString(inch, 0.5 * inch, f"{TITLE} — {SUBTITLE}")
        canvas.drawRightString(LETTER[0] - inch, 0.5 * inch, f"Page {doc.page}")
        canvas.restoreState()

    def afterFlowable(self, flowable):
        if not hasattr(flowable, "_toc_level"):
            return
        text = flowable.getPlainText()
        self.notify("TOCEntry", (flowable._toc_level, text, self.page))


def build_pdf(out_path):
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["Normal"], fontName="Helvetica",
                          fontSize=10, leading=14, spaceAfter=6)
    h_styles = {
        1: ParagraphStyle("H1", parent=styles["Heading1"], fontName="Helvetica-Bold",
                          fontSize=20, leading=24, spaceBefore=6, spaceAfter=10,
                          textColor=colors.HexColor("#212529")),
        2: ParagraphStyle("H2", parent=styles["Heading2"], fontName="Helvetica-Bold",
                          fontSize=15, leading=19, spaceBefore=12, spaceAfter=6,
                          textColor=colors.HexColor("#0d6efd")),
        3: ParagraphStyle("H3", parent=styles["Heading3"], fontName="Helvetica-Bold",
                          fontSize=12, leading=16, spaceBefore=8, spaceAfter=4,
                          textColor=colors.HexColor("#343a40")),
        4: ParagraphStyle("H4", parent=styles["Heading4"], fontName="Helvetica-BoldOblique",
                          fontSize=11, leading=14, spaceBefore=6, spaceAfter=3),
    }
    code_style = ParagraphStyle("Code", fontName="Courier", fontSize=7.8, leading=9.6,
                                textColor=colors.HexColor("#212529"))
    callout_style = ParagraphStyle("Callout", parent=body, leftIndent=14,
                                   textColor=colors.HexColor("#495057"),
                                   borderColor=colors.HexColor("#0d6efd"),
                                   borderWidth=0, spaceBefore=4, spaceAfter=4)
    li_style = ParagraphStyle("LI", parent=body, spaceAfter=2)

    story = []

    # Title page
    story.append(Spacer(1, 2.2 * inch))
    story.append(Paragraph(TITLE, ParagraphStyle("T", parent=body, fontName="Helvetica-Bold",
                                                  fontSize=40, alignment=TA_CENTER, leading=46)))
    story.append(Spacer(1, 0.2 * inch))
    story.append(Paragraph(SUBTITLE, ParagraphStyle("S", parent=body, fontSize=20,
                 alignment=TA_CENTER, textColor=colors.HexColor("#495057"))))
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(VERSION, ParagraphStyle("V", parent=body, fontSize=11,
                 alignment=TA_CENTER, textColor=colors.grey)))
    story.append(Paragraph(f"Generated {_dt.date.today().isoformat()}",
                 ParagraphStyle("D", parent=body, fontSize=10, alignment=TA_CENTER,
                                textColor=colors.grey)))
    story.append(PageBreak())

    # TOC
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle("TOC1", fontName="Helvetica-Bold", fontSize=11, leading=16),
        ParagraphStyle("TOC2", fontName="Helvetica", fontSize=9.5, leading=13, leftIndent=18),
    ]
    story.append(Paragraph("Table of Contents",
                 ParagraphStyle("TOCH", parent=h_styles[1], alignment=TA_LEFT)))
    story.append(toc)
    story.append(PageBreak())

    def render_block(node, container):
        if isinstance(node, NavigableString):
            txt = str(node).strip()
            if txt:
                container.append(Paragraph(_html.escape(txt), body))
            return
        if not isinstance(node, Tag):
            return
        name = node.name.lower()
        if re.fullmatch(r"h[1-6]", name):
            level = min(int(name[1]), 4)
            style = h_styles[level]
            para = Paragraph(_inline_to_rl(node), style)
            if level <= 2:
                para._toc_level = level - 1  # H1->0, H2->1
            container.append(para)
        elif name == "p":
            container.append(Paragraph(_inline_to_rl(node) or "&nbsp;", body))
        elif name in ("ul", "ol"):
            container.append(_build_list(node))
        elif name == "table":
            t = _build_table(node)
            if t:
                container.append(t)
                container.append(Spacer(1, 6))
        elif name == "pre":
            code = node.find("code")
            text = code.get_text() if code else node.get_text()
            text = wrap_pre_lines(text.rstrip("\n"))
            box = Table([[Preformatted(text, code_style)]], colWidths=[6.5 * inch])
            box.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F4F4F4")),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            container.append(box)
            container.append(Spacer(1, 6))
        elif name == "blockquote":
            inner = node.find_all("p")
            text = "<br/>".join(_inline_to_rl(p) for p in inner) if inner else _inline_to_rl(node)
            box = Table([[Paragraph(text, callout_style)]], colWidths=[6.5 * inch])
            box.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EAF2FB")),
                ("LINEBEFORE", (0, 0), (0, -1), 3, colors.HexColor("#0d6efd")),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            container.append(box)
            container.append(Spacer(1, 6))
        elif name == "hr":
            container.append(HRFlowable(width="100%", thickness=0.5,
                             color=colors.HexColor("#CCCCCC"), spaceBefore=6, spaceAfter=6))
        else:
            for child in node.children:
                render_block(child, container)

    def _build_list(node, depth=0):
        ordered = node.name.lower() == "ol"
        items = []
        for li in node.find_all("li", recursive=False):
            parts = []
            inline = "".join(_inline_to_rl(c) for c in li.children
                             if not (isinstance(c, Tag) and c.name in ("ul", "ol")))
            if inline.strip():
                parts.append(Paragraph(inline, li_style))
            for sub in li.find_all(["ul", "ol"], recursive=False):
                parts.append(_build_list(sub, depth + 1))
            items.append(ListItem(parts if len(parts) > 1 else (parts[0] if parts else Paragraph("", li_style)),
                                  leftIndent=18 + depth * 12))
        return ListFlowable(
            items,
            bulletType="1" if ordered else "bullet",
            start="1" if ordered else None,
            bulletColor=colors.HexColor("#0d6efd"),
            leftIndent=18 + depth * 12,
        )

    def _build_table(node):
        rows = node.find_all("tr")
        if not rows:
            return None
        ncols = max(len(r.find_all(["td", "th"])) for r in rows)
        data = []
        header_rows = 0
        cell_para = ParagraphStyle("Cell", parent=body, fontSize=8.5, leading=11, spaceAfter=0)
        head_para = ParagraphStyle("CellH", parent=cell_para, fontName="Helvetica-Bold",
                                   textColor=colors.white)
        for tr in rows:
            cells = tr.find_all(["td", "th"])
            is_header = all(c.name == "th" for c in cells) and cells
            if is_header:
                header_rows += 1
            row = []
            for c_idx in range(ncols):
                if c_idx < len(cells):
                    style = head_para if cells[c_idx].name == "th" else cell_para
                    row.append(Paragraph(_inline_to_rl(cells[c_idx]) or "&nbsp;", style))
                else:
                    row.append(Paragraph("", cell_para))
            data.append(row)
        col_w = (6.5 * inch) / ncols
        t = Table(data, colWidths=[col_w] * ncols, repeatRows=header_rows)
        ts = [
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#CCCCCC")),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("ROWBACKGROUNDS", (0, header_rows), (-1, -1),
             [colors.white, colors.HexColor("#F7F9FC")]),
        ]
        if header_rows:
            ts.append(("BACKGROUND", (0, 0), (-1, header_rows - 1), colors.HexColor("#0d6efd")))
        t.setStyle(TableStyle(ts))
        return t

    for rel in DOC_FILES:
        path = os.path.join(HERE, rel)
        if not os.path.exists(path):
            continue
        with open(path, encoding="utf-8") as f:
            soup = md_to_soup(f.read())
        for block in list(soup.children):
            render_block(block, story)
        story.append(PageBreak())

    doc = ErpDocTemplate(out_path, pagesize=LETTER,
                         leftMargin=inch, rightMargin=inch,
                         topMargin=inch, bottomMargin=0.9 * inch,
                         title=f"{TITLE} {SUBTITLE}")
    doc.multiBuild(story)
    print(f"  PDF written: {out_path}")


def main():
    docx_out = os.path.join(HERE, OUT_BASE + ".docx")
    pdf_out = os.path.join(HERE, OUT_BASE + ".pdf")
    print("Building ChuA.ERP user documentation...")
    build_docx(docx_out)
    build_pdf(pdf_out)
    print("Done.")


if __name__ == "__main__":
    sys.exit(main())
